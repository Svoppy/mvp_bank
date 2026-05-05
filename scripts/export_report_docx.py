from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLACK = (0, 0, 0)
WHITE = (255, 255, 255)


@dataclass
class Box:
    x: int
    y: int
    w: int
    h: int

    @property
    def left(self) -> int:
        return self.x

    @property
    def right(self) -> int:
        return self.x + self.w

    @property
    def top(self) -> int:
        return self.y

    @property
    def bottom(self) -> int:
        return self.y + self.h

    @property
    def center(self) -> tuple[int, int]:
        return (self.x + self.w // 2, self.y + self.h // 2)


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_name = "timesbd.ttf" if bold else "times.ttf"
    font_path = Path("C:/Windows/Fonts") / font_name
    if font_path.is_file():
        return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


FONT_30 = load_font(30)
FONT_30_BOLD = load_font(30, bold=True)
FONT_26 = load_font(26)
FONT_26_BOLD = load_font(26, bold=True)
FONT_24 = load_font(24)
FONT_24_BOLD = load_font(24, bold=True)
FONT_22 = load_font(22)
FONT_22_BOLD = load_font(22, bold=True)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def multiline_size(draw: ImageDraw.ImageDraw, lines: list[str], font: ImageFont.ImageFont, spacing: int = 6) -> tuple[int, int]:
    widths = [draw.textbbox((0, 0), line, font=font)[2] for line in lines] or [0]
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3]
    height = line_height * len(lines) + spacing * max(len(lines) - 1, 0)
    return max(widths), height


def draw_box(draw: ImageDraw.ImageDraw, box: Box, text: str, *, font: ImageFont.ImageFont = FONT_24, bold: bool = False) -> None:
    draw.rectangle([box.left, box.top, box.right, box.bottom], outline=BLACK, width=3, fill=WHITE)
    text_font = FONT_24_BOLD if bold else font
    lines = wrap_text(draw, text, text_font, box.w - 24)
    _, text_height = multiline_size(draw, lines, text_font)
    current_y = box.top + (box.h - text_height) // 2
    for line in lines:
        text_width = draw.textbbox((0, 0), line, font=text_font)[2]
        draw.text((box.left + (box.w - text_width) // 2, current_y), line, fill=BLACK, font=text_font)
        current_y += draw.textbbox((0, 0), "Ag", font=text_font)[3] + 6


def draw_diamond(draw: ImageDraw.ImageDraw, box: Box, text: str) -> None:
    points = [
        (box.center[0], box.top),
        (box.right, box.center[1]),
        (box.center[0], box.bottom),
        (box.left, box.center[1]),
    ]
    draw.polygon(points, outline=BLACK, fill=WHITE, width=3)
    lines = wrap_text(draw, text, FONT_22_BOLD, int(box.w * 0.58))
    _, text_height = multiline_size(draw, lines, FONT_22_BOLD)
    current_y = box.top + (box.h - text_height) // 2
    for line in lines:
        text_width = draw.textbbox((0, 0), line, font=FONT_22_BOLD)[2]
        draw.text((box.left + (box.w - text_width) // 2, current_y), line, fill=BLACK, font=FONT_22_BOLD)
        current_y += draw.textbbox((0, 0), "Ag", font=FONT_22_BOLD)[3] + 4


def draw_group(
    draw: ImageDraw.ImageDraw,
    box: Box,
    title: str,
    nodes: list[tuple[str, str]],
    *,
    single_width: int | None = None,
    top_padding: int = 90,
    gap: int = 28,
) -> dict[str, Box]:
    draw.rectangle([box.left, box.top, box.right, box.bottom], outline=BLACK, width=4, fill=WHITE)
    title_lines = wrap_text(draw, title, FONT_26_BOLD, box.w - 40)
    title_y = box.top + 16
    for line in title_lines:
        width = draw.textbbox((0, 0), line, font=FONT_26_BOLD)[2]
        draw.text((box.left + (box.w - width) // 2, title_y), line, fill=BLACK, font=FONT_26_BOLD)
        title_y += draw.textbbox((0, 0), "Ag", font=FONT_26_BOLD)[3] + 4

    node_map: dict[str, Box] = {}
    current_y = box.top + top_padding
    width = single_width or box.w - 60
    x = box.left + (box.w - width) // 2
    for node_id, label in nodes:
        height = 86 if len(label) < 30 else 110
        node_box = Box(x, current_y, width, height)
        draw_box(draw, node_box, label)
        node_map[node_id] = node_box
        current_y += height + gap
    return node_map


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    *,
    label: str | None = None,
    label_offset: tuple[int, int] = (0, 0),
) -> None:
    draw.line(points, fill=BLACK, width=3)
    if len(points) < 2:
        return

    (x1, y1), (x2, y2) = points[-2], points[-1]
    dx = x2 - x1
    dy = y2 - y1
    if dx == 0 and dy == 0:
        return

    length = (dx**2 + dy**2) ** 0.5
    ux = dx / length
    uy = dy / length
    arrow_size = 14
    px = -uy
    py = ux
    tip = (x2, y2)
    left = (int(x2 - ux * arrow_size + px * arrow_size * 0.6), int(y2 - uy * arrow_size + py * arrow_size * 0.6))
    right = (int(x2 - ux * arrow_size - px * arrow_size * 0.6), int(y2 - uy * arrow_size - py * arrow_size * 0.6))
    draw.polygon([tip, left, right], outline=BLACK, fill=BLACK)

    if label:
        mid_index = max((len(points) - 1) // 2, 0)
        lx = points[mid_index][0] + label_offset[0]
        ly = points[mid_index][1] + label_offset[1]
        label_lines = wrap_text(draw, label, FONT_22, 220)
        width, height = multiline_size(draw, label_lines, FONT_22, spacing=2)
        padding = 8
        rect = [lx - padding, ly - padding, lx + width + padding, ly + height + padding]
        draw.rectangle(rect, fill=WHITE)
        current_y = ly
        for line in label_lines:
            draw.text((lx, current_y), line, fill=BLACK, font=FONT_22)
            current_y += draw.textbbox((0, 0), "Ag", font=FONT_22)[3] + 2


def render_diagram_1(path: Path) -> None:
    image = Image.new("RGB", (1900, 1180), WHITE)
    draw = ImageDraw.Draw(image)

    groups = {
        "U": Box(40, 90, 290, 960),
        "T": Box(380, 360, 330, 240),
        "APP": Box(760, 60, 560, 1030),
        "DATA": Box(1380, 140, 420, 900),
    }

    nodes_u = draw_group(
        draw,
        groups["U"],
        "Недоверенная зона",
        [("C", "CLIENT"), ("M", "MANAGER"), ("A", "ADMIN")],
        single_width=180,
        top_padding=170,
        gap=80,
    )
    nodes_t = draw_group(draw, groups["T"], "Граница доверия: HTTPS", [("H", "HTTPS endpoint 127.0.0.1:7443")], single_width=250, top_padding=110)
    nodes_app = draw_group(
        draw,
        groups["APP"],
        "Приложение Django / Django Ninja",
        [
            ("API", "API routers: auth / loans / audit"),
            ("AUTH", "JWTAuth + decode_token()"),
            ("LOANS", "Loans API"),
            ("AUDIT", "Audit API"),
            ("NET", "get_client_ip()"),
            ("DOCS", "store_loan_document()"),
            ("CSV", "stream_loans_csv_response()"),
        ],
        single_width=360,
        top_padding=130,
        gap=26,
    )
    nodes_data = draw_group(
        draw,
        groups["DATA"],
        "Доверенная зона хранения",
        [("DB", "SQLite / PostgreSQL"), ("FS", "MEDIA_ROOT/loan_documents"), ("LOG", "AuditLog + app logs")],
        single_width=300,
        top_padding=160,
        gap=80,
    )

    draw_arrow(
        draw,
        [(nodes_u["C"].right, nodes_u["C"].center[1]), (nodes_t["H"].left, nodes_t["H"].center[1] - 32)],
        label="JSON / multipart",
        label_offset=(30, -40),
    )
    draw_arrow(
        draw,
        [(nodes_u["M"].right, nodes_u["M"].center[1]), (nodes_t["H"].left, nodes_t["H"].center[1])],
        label="JSON",
        label_offset=(40, -18),
    )
    draw_arrow(
        draw,
        [(nodes_u["A"].right, nodes_u["A"].center[1]), (nodes_t["H"].left, nodes_t["H"].center[1] + 32)],
        label="JSON",
        label_offset=(40, 10),
    )
    draw_arrow(draw, [(nodes_t["H"].right, nodes_t["H"].center[1]), (nodes_app["API"].left, nodes_app["API"].center[1])])
    draw_arrow(draw, [(nodes_app["API"].center[0], nodes_app["API"].bottom), (nodes_app["AUTH"].center[0], nodes_app["AUTH"].top)])
    draw_arrow(draw, [(nodes_app["API"].center[0], nodes_app["API"].bottom), (nodes_app["LOANS"].center[0], nodes_app["LOANS"].top)])
    draw_arrow(draw, [(nodes_app["API"].center[0], nodes_app["API"].bottom), (nodes_app["AUDIT"].center[0], nodes_app["AUDIT"].top)])
    draw_arrow(draw, [(nodes_app["API"].center[0], nodes_app["API"].bottom), (nodes_app["NET"].center[0], nodes_app["NET"].top)])
    draw_arrow(
        draw,
        [(nodes_app["LOANS"].right, nodes_app["LOANS"].center[1]), (nodes_data["DB"].left - 30, nodes_app["LOANS"].center[1]), (nodes_data["DB"].left - 30, nodes_data["DB"].center[1]), (nodes_data["DB"].left, nodes_data["DB"].center[1])],
    )
    draw_arrow(draw, [(nodes_app["LOANS"].center[0], nodes_app["LOANS"].bottom), (nodes_app["DOCS"].center[0], nodes_app["DOCS"].top)])
    draw_arrow(draw, [(nodes_app["LOANS"].center[0], nodes_app["LOANS"].bottom), (nodes_app["CSV"].center[0], nodes_app["CSV"].top)])
    draw_arrow(
        draw,
        [(nodes_app["DOCS"].right, nodes_app["DOCS"].center[1]), (nodes_data["FS"].left - 30, nodes_app["DOCS"].center[1]), (nodes_data["FS"].left - 30, nodes_data["FS"].center[1]), (nodes_data["FS"].left, nodes_data["FS"].center[1])],
    )
    draw_arrow(
        draw,
        [(nodes_app["AUDIT"].right, nodes_app["AUDIT"].center[1]), (nodes_data["DB"].left - 65, nodes_app["AUDIT"].center[1]), (nodes_data["DB"].left - 65, nodes_data["DB"].center[1]), (nodes_data["DB"].left, nodes_data["DB"].center[1])],
    )
    draw_arrow(
        draw,
        [(nodes_app["API"].right, nodes_app["API"].center[1]), (nodes_data["LOG"].left - 40, nodes_app["API"].center[1]), (nodes_data["LOG"].left - 40, nodes_data["LOG"].center[1]), (nodes_data["LOG"].left, nodes_data["LOG"].center[1])],
    )
    draw_arrow(
        draw,
        [(nodes_app["CSV"].left, nodes_app["CSV"].center[1]), (groups["T"].right + 40, nodes_app["CSV"].center[1]), (groups["T"].right + 40, nodes_t["H"].center[1]), (nodes_t["H"].right, nodes_t["H"].center[1])],
    )
    image.save(path)


def render_diagram_2(path: Path) -> None:
    image = Image.new("RGB", (1600, 2320), WHITE)
    draw = ImageDraw.Draw(image)

    boxes = {
        "S": Box(540, 40, 420, 74),
        "T": Box(540, 155, 420, 74),
        "A1": Box(540, 270, 420, 86),
        "D1": Box(520, 410, 460, 160),
        "E1": Box(1130, 460, 180, 64),
        "C1": Box(500, 640, 500, 86),
        "U1": Box(480, 770, 540, 96),
        "D2": Box(500, 930, 500, 170),
        "E2": Box(1130, 980, 180, 64),
        "V1": Box(440, 1170, 620, 96),
        "W1": Box(390, 1310, 720, 120),
        "M1": Box(520, 1480, 460, 74),
        "L1": Box(520, 1590, 460, 74),
        "P1": Box(450, 1700, 600, 96),
        "D3": Box(520, 1850, 460, 160),
        "E3": Box(1130, 1900, 180, 64),
        "Q1": Box(430, 2080, 640, 96),
        "R1": Box(520, 2210, 460, 74),
        "F": Box(600, 2310, 300, 0),
    }

    draw_box(draw, boxes["S"], "Клиент логинится")
    draw_box(draw, boxes["T"], "JWT access token выдан")
    draw_box(draw, boxes["A1"], "POST /api/loans/apply")
    draw_diamond(draw, boxes["D1"], "Роль = CLIENT?")
    draw_box(draw, boxes["E1"], "403")
    draw_box(draw, boxes["C1"], "Создать CreditApplication")
    draw_box(draw, boxes["U1"], "POST /api/loans/{loan_id}/documents")
    draw_diamond(draw, boxes["D2"], "Заявка принадлежит клиенту?")
    draw_box(draw, boxes["E2"], "404")
    draw_box(draw, boxes["V1"], "Проверка размера / magic bytes / safe path")
    draw_box(draw, boxes["W1"], "Сохранить файл потоково и записать LoanDocument")
    draw_box(draw, boxes["M1"], "Менеджер логинится")
    draw_box(draw, boxes["L1"], "GET /api/loans/")
    draw_box(draw, boxes["P1"], "PATCH /api/loans/{loan_id}/decision")
    draw_diamond(draw, boxes["D3"], "Роль = MANAGER?")
    draw_box(draw, boxes["E3"], "403")
    draw_box(draw, boxes["Q1"], "select_for_update() и смена статуса")
    draw_box(draw, boxes["R1"], "Запись в AuditLog")
    draw_box(draw, Box(610, 2225, 280, 74), "Готово")

    draw_arrow(draw, [(boxes["S"].center[0], boxes["S"].bottom), (boxes["T"].center[0], boxes["T"].top)])
    draw_arrow(draw, [(boxes["T"].center[0], boxes["T"].bottom), (boxes["A1"].center[0], boxes["A1"].top)])
    draw_arrow(draw, [(boxes["A1"].center[0], boxes["A1"].bottom), (boxes["D1"].center[0], boxes["D1"].top)])
    draw_arrow(draw, [(boxes["D1"].right, boxes["D1"].center[1]), (boxes["E1"].left, boxes["E1"].center[1])], label="нет", label_offset=(30, -20))
    draw_arrow(draw, [(boxes["D1"].center[0], boxes["D1"].bottom), (boxes["C1"].center[0], boxes["C1"].top)], label="да", label_offset=(20, -15))
    draw_arrow(draw, [(boxes["C1"].center[0], boxes["C1"].bottom), (boxes["U1"].center[0], boxes["U1"].top)])
    draw_arrow(draw, [(boxes["U1"].center[0], boxes["U1"].bottom), (boxes["D2"].center[0], boxes["D2"].top)])
    draw_arrow(draw, [(boxes["D2"].right, boxes["D2"].center[1]), (boxes["E2"].left, boxes["E2"].center[1])], label="нет", label_offset=(30, -20))
    draw_arrow(draw, [(boxes["D2"].center[0], boxes["D2"].bottom), (boxes["V1"].center[0], boxes["V1"].top)], label="да", label_offset=(20, -15))
    draw_arrow(draw, [(boxes["V1"].center[0], boxes["V1"].bottom), (boxes["W1"].center[0], boxes["W1"].top)])
    draw_arrow(draw, [(boxes["W1"].center[0], boxes["W1"].bottom), (boxes["M1"].center[0], boxes["M1"].top)])
    draw_arrow(draw, [(boxes["M1"].center[0], boxes["M1"].bottom), (boxes["L1"].center[0], boxes["L1"].top)])
    draw_arrow(draw, [(boxes["L1"].center[0], boxes["L1"].bottom), (boxes["P1"].center[0], boxes["P1"].top)])
    draw_arrow(draw, [(boxes["P1"].center[0], boxes["P1"].bottom), (boxes["D3"].center[0], boxes["D3"].top)])
    draw_arrow(draw, [(boxes["D3"].right, boxes["D3"].center[1]), (boxes["E3"].left, boxes["E3"].center[1])], label="нет", label_offset=(30, -20))
    draw_arrow(draw, [(boxes["D3"].center[0], boxes["D3"].bottom), (boxes["Q1"].center[0], boxes["Q1"].top)], label="да", label_offset=(20, -15))
    draw_arrow(draw, [(boxes["Q1"].center[0], boxes["Q1"].bottom), (boxes["R1"].center[0], boxes["R1"].top)])
    draw_arrow(draw, [(boxes["R1"].center[0], boxes["R1"].bottom), (750, 2225)])
    image.crop((0, 0, 1600, 2310)).save(path)


def render_diagram_3(path: Path) -> None:
    image = Image.new("RGB", (1820, 980), WHITE)
    draw = ImageDraw.Draw(image)

    groups = {
        "B1": Box(40, 140, 340, 600),
        "B2": Box(430, 140, 360, 640),
        "B3": Box(860, 110, 420, 760),
        "B4": Box(1360, 220, 360, 500),
    }

    nodes_b1 = draw_group(draw, groups["B1"], "Boundary 1: Интернет / клиент", [("REQ", "HTTP JSON, query params, path params, multipart file")], single_width=260, top_padding=170)
    nodes_b2 = draw_group(draw, groups["B2"], "Boundary 2: API and validation", [("SCHEMA", "Pydantic schemas + query constraints"), ("AUTHZ", "JWTAuth + role checks + object checks")], single_width=280, top_padding=170, gap=80)
    nodes_b3 = draw_group(draw, groups["B3"], "Boundary 3: Trusted execution", [("ORM", "Django ORM"), ("FILES", "Safe file storage"), ("CSV", "Streaming CSV serializer"), ("LOGS", "Audit logging")], single_width=280, top_padding=150, gap=46)
    nodes_b4 = draw_group(draw, groups["B4"], "Boundary 4: Data stores", [("DB2", "Database"), ("FS2", "File system")], single_width=220, top_padding=170, gap=100)

    draw_arrow(draw, [(nodes_b1["REQ"].right, nodes_b1["REQ"].center[1] - 25), (nodes_b2["SCHEMA"].left, nodes_b2["SCHEMA"].center[1])])
    draw_arrow(draw, [(nodes_b1["REQ"].right, nodes_b1["REQ"].center[1] + 25), (nodes_b2["AUTHZ"].left, nodes_b2["AUTHZ"].center[1])])
    draw_arrow(draw, [(nodes_b2["SCHEMA"].right, nodes_b2["SCHEMA"].center[1]), (nodes_b3["ORM"].left, nodes_b3["ORM"].center[1])])
    draw_arrow(draw, [(nodes_b2["AUTHZ"].right, nodes_b2["AUTHZ"].center[1]), (nodes_b3["ORM"].left, nodes_b3["ORM"].center[1] + 42)])
    draw_arrow(draw, [(nodes_b3["ORM"].right, nodes_b3["ORM"].center[1]), (nodes_b4["DB2"].left, nodes_b4["DB2"].center[1])])
    draw_arrow(draw, [(nodes_b2["SCHEMA"].right, nodes_b2["SCHEMA"].bottom), (nodes_b3["FILES"].left, nodes_b3["FILES"].center[1])])
    draw_arrow(draw, [(nodes_b3["FILES"].right, nodes_b3["FILES"].center[1]), (nodes_b4["FS2"].left, nodes_b4["FS2"].center[1])])
    draw_arrow(draw, [(nodes_b3["ORM"].center[0], nodes_b3["ORM"].bottom), (nodes_b3["CSV"].center[0], nodes_b3["CSV"].top)])
    draw_arrow(draw, [(nodes_b3["CSV"].left, nodes_b3["CSV"].center[1]), (nodes_b1["REQ"].right - 40, nodes_b3["CSV"].center[1]), (nodes_b1["REQ"].right - 40, nodes_b1["REQ"].bottom + 40), (nodes_b1["REQ"].center[0], nodes_b1["REQ"].bottom)], label_offset=(0, 0))
    draw_arrow(draw, [(nodes_b2["AUTHZ"].right, nodes_b2["AUTHZ"].center[1]), (nodes_b3["LOGS"].left, nodes_b3["LOGS"].center[1])])
    draw_arrow(draw, [(nodes_b3["LOGS"].right, nodes_b3["LOGS"].center[1]), (nodes_b4["DB2"].left, nodes_b4["DB2"].center[1] + 50)])
    image.save(path)


def render_mermaid_block(block_index: int, path: Path) -> None:
    if block_index == 1:
        render_diagram_1(path)
    elif block_index == 2:
        render_diagram_2(path)
    elif block_index == 3:
        render_diagram_3(path)
    else:
        image = Image.new("RGB", (1200, 240), WHITE)
        draw = ImageDraw.Draw(image)
        draw.rectangle([20, 20, 1180, 220], outline=BLACK, width=3)
        draw.text((40, 90), "Mermaid diagram omitted", fill=BLACK, font=FONT_30_BOLD)
        image.save(path)


def normalize_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def add_run(paragraph, text: str, *, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_inline_markdown(paragraph, text: str) -> None:
    normalized = normalize_links(text)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", normalized)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_run(paragraph, part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            add_run(paragraph, part[1:-1])
        else:
            add_run(paragraph, part)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_markdown(p, text)


def set_paragraph_format(paragraph, *, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.bold = bold if run.bold is None else run.bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)


def parse_table_line(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    allowed = set("|:- ")
    return all(char in allowed for char in stripped)


def apply_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"):
        style_obj = document.styles[style_name]
        style_obj.font.name = "Times New Roman"
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style_obj.font.size = Pt(12)
        style_obj.font.bold = True
        style_obj.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(document: Document, line: str) -> None:
    match = re.match(r"^(#{1,6})\s+(.*)$", line)
    if not match:
        return
    level = len(match.group(1))
    text = normalize_links(match.group(2))
    paragraph = document.add_paragraph()
    add_run(paragraph, text, bold=True)
    set_paragraph_format(paragraph, bold=True)
    paragraph.style = f"Heading {min(level, 6)}"


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    add_inline_markdown(paragraph, text)
    set_paragraph_format(paragraph)


def add_list_item(document: Document, text: str, *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    add_inline_markdown(paragraph, text)
    set_paragraph_format(paragraph)


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        add_run(p, line)
        if index < len(lines) - 1:
            p.add_run("\n")
    document.add_paragraph("")


def add_table(document: Document, lines: list[str]) -> None:
    rows = [parse_table_line(line) for line in lines if line.strip()]
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, cell_text in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline_markdown(p, cell_text)
        set_paragraph_format(p, bold=True)

    for row_data in body:
        row = table.add_row().cells
        for index, cell_text in enumerate(row_data):
            set_cell_text(row[index], cell_text)

    document.add_paragraph("")


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    assets_dir = output_path.parent / "docx_assets"
    assets_dir.mkdir(parents=True, exist_ok=True)

    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    apply_normal_style(document)
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    mermaid_index = 0
    diagram_widths_cm = {
        1: 15.8,
        2: 10.5,
        3: 15.8,
    }
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            language = line[3:].strip()
            block_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block_lines.append(lines[i])
                i += 1
            i += 1
            if language == "mermaid":
                mermaid_index += 1
                image_path = assets_dir / f"diagram_{mermaid_index}.png"
                render_mermaid_block(mermaid_index, image_path)
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Cm(diagram_widths_cm.get(mermaid_index, 15.8)))
                paragraph.paragraph_format.space_after = Pt(6)
            else:
                add_code_block(document, block_lines)
            continue

        if re.match(r"^(#{1,6})\s+", line):
            add_heading(document, line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        if re.match(r"^\s*-\s+", line):
            add_list_item(document, re.sub(r"^\s*-\s+", "", line).strip(), numbered=False)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            add_list_item(document, re.sub(r"^\s*\d+\.\s+", "", line).strip(), numbered=True)
            i += 1
            continue

        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            current = lines[i]
            if not current.strip():
                break
            if current.startswith("```") or re.match(r"^(#{1,6})\s+", current) or re.match(r"^\s*-\s+", current) or re.match(r"^\s*\d+\.\s+", current):
                break
            if current.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                break
            paragraph_lines.append(current.strip())
            i += 1
        add_paragraph(document, " ".join(paragraph_lines))

    document.save(output_path)


def main(argv: Iterable[str]) -> int:
    args = list(argv)
    if len(args) != 2:
        print("Usage: python scripts/export_report_docx.py <input.md> <output.docx>")
        return 1

    input_path = Path(args[0]).resolve()
    output_path = Path(args[1]).resolve()
    convert_markdown_to_docx(input_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
